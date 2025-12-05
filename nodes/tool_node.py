# ---------------------------------------------------------
# Tool Node 1: Word 파일 생성
# ---------------------------------------------------------

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def tool_node1(state, llm, tools):
    """
    integrator_node에서 생성한 통합 보고서를 Word 파일로 저장하는 노드입니다.

    Args:
        state: 현재 State 객체 (answer 필드를 포함)
        llm: LLM 인스턴스 (사용하지 않지만 호환성을 위해 유지)
        tools: 사용 가능한 툴 리스트 (사용하지 않지만 호환성을 위해 유지)

    Returns:
        dict: 업데이트된 State (word_file_path 필드 추가)
    """
    print("=== Word 파일 생성 시작 ===")
    
    # integrator_node에서 생성한 통합 보고서 가져오기
    report_content = state.get("answer", "")
    
    if not report_content or report_content.strip() == "":
        print("경고: 보고서 내용이 없습니다.")
        return {"word_file_path": None}
    
    # 출력 폴더 생성
    output_folder = "outputs"
    os.makedirs(output_folder, exist_ok=True)
    
    # 파일명 생성 (타임스탬프 포함)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"asset_analysis_{timestamp}.docx"
    file_path = os.path.join(output_folder, filename)
    
    # Word 문서 생성
    doc = Document()
    
    # 제목 스타일 설정
    title = doc.add_heading('Asset Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목 (타임스탬프)
    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(10)
    subtitle_format.color = RGBColor(128, 128, 128)
    
    # 구분선 추가
    doc.add_paragraph("=" * 80)
    
    # 보고서 내용 파싱 및 추가
    lines = report_content.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # 빈 줄 추가
            if current_paragraph:
                doc.add_paragraph()
            continue
        
        # 섹션 헤더 감지 (===== 또는 [SECTION] 형식)
        if line.startswith('=' * 10) or (line.startswith('[') and line.endswith(']')):
            # 섹션 제목
            if line.startswith('[') and line.endswith(']'):
                heading = doc.add_heading(line[1:-1], level=1)
                heading_format = heading.runs[0].font
                heading_format.bold = True
                heading_format.size = Pt(14)
                heading_format.color = RGBColor(0, 0, 139)  # Dark blue
            else:
                # 구분선은 그대로 추가
                doc.add_paragraph(line)
        elif line.startswith('**') and line.endswith('**'):
            # 볼드 텍스트 (서브섹션)
            para = doc.add_paragraph()
            run = para.add_run(line.replace('**', ''))
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('- ') or line.startswith('* '):
            # 리스트 항목
            para = doc.add_paragraph(line[2:], style='List Bullet')
            para_format = para.runs[0].font
            para_format.size = Pt(11)
        elif any(keyword in line for keyword in ['[', ']', '**']):
            # 마크다운 형식이 포함된 일반 텍스트
            para = doc.add_paragraph()
            # 간단한 마크다운 파싱
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.bold = True
            para_format = para.runs[0].font
            para_format.size = Pt(11)
        else:
            # 일반 텍스트
            para = doc.add_paragraph(line)
            para_format = para.runs[0].font
            para_format.size = Pt(11)
            current_paragraph = para
    
    # 문서 저장
    doc.save(file_path)
    
    print(f"Word 파일 생성 완료: {file_path}")
    print(f"파일 크기: {os.path.getsize(file_path)} bytes")
    
    return {"word_file_path": file_path}

